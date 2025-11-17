"""Document processing service for various file types."""
import os
import shutil
from pathlib import Path
from typing import List, Dict, Any, Optional
from uuid import UUID
import PyPDF2
import pdfplumber
import pandas as pd
from docx import Document as DocxDocument

from app.config import settings
from app.utils.logger import logger


class DocumentProcessor:
    """Service for processing different document types."""

    def __init__(self):
        """Initialize document processor."""
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP

    def process_file(
        self,
        file_path: Path,
        file_type: str,
        password: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """
        Process file and extract text chunks.

        Args:
            file_path: Path to file
            file_type: Type of file (pdf, xlsx, csv, docx)
            password: Optional password for protected files

        Returns:
            List of text chunks with metadata
        """
        try:
            if file_type == "pdf":
                return self.process_pdf(file_path, password)
            elif file_type in ["xlsx", "xls"]:
                return self.process_excel(file_path)
            elif file_type == "csv":
                return self.process_csv(file_path)
            elif file_type == "docx":
                return self.process_docx(file_path)
            else:
                logger.error(f"Unsupported file type: {file_type}")
                return []

        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            return []

    def process_pdf(
        self,
        file_path: Path,
        password: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """
        Process PDF file and extract text.

        Args:
            file_path: Path to PDF file
            password: Optional password for protected PDFs

        Returns:
            List of text chunks with page numbers
        """
        chunks = []

        try:
            # Try with pdfplumber first (better for tables)
            with pdfplumber.open(file_path, password=password) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text()

                    if text:
                        # Extract tables if any
                        tables = page.extract_tables()
                        if tables:
                            table_text = self._format_tables(tables)
                            text += "\n\n" + table_text

                        # Create chunks from page text
                        page_chunks = self._create_chunks(text, page_number=page_num)
                        chunks.extend(page_chunks)

            logger.info(f"Processed PDF: {file_path.name}, extracted {len(chunks)} chunks")
            return chunks

        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {str(e)}")

            try:
                # Fallback to PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)

                    # Decrypt if password protected
                    if pdf_reader.is_encrypted:
                        if password:
                            pdf_reader.decrypt(password)
                        else:
                            raise ValueError("PDF is password protected but no password provided")

                    for page_num, page in enumerate(pdf_reader.pages, start=1):
                        text = page.extract_text()

                        if text:
                            page_chunks = self._create_chunks(text, page_number=page_num)
                            chunks.extend(page_chunks)

                logger.info(f"Processed PDF with PyPDF2: {file_path.name}, extracted {len(chunks)} chunks")
                return chunks

            except Exception as fallback_error:
                logger.error(f"Both PDF processors failed: {str(fallback_error)}")
                return []

    def process_excel(self, file_path: Path) -> List[Dict[str, Any]]:
        """
        Process Excel file and extract data.

        Args:
            file_path: Path to Excel file

        Returns:
            List of text chunks from sheets
        """
        chunks = []

        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)

                # Convert DataFrame to text
                text = f"Sheet: {sheet_name}\n\n"
                text += df.to_string(index=False)

                # Create chunks with sheet metadata
                sheet_chunks = self._create_chunks(
                    text,
                    metadata={"sheet_name": sheet_name}
                )
                chunks.extend(sheet_chunks)

            logger.info(f"Processed Excel: {file_path.name}, extracted {len(chunks)} chunks")
            return chunks

        except Exception as e:
            logger.error(f"Error processing Excel file: {str(e)}")
            return []

    def process_csv(self, file_path: Path) -> List[Dict[str, Any]]:
        """
        Process CSV file and extract data.

        Args:
            file_path: Path to CSV file

        Returns:
            List of text chunks
        """
        try:
            df = pd.read_csv(file_path)

            # Convert DataFrame to text
            text = df.to_string(index=False)

            # Create chunks
            chunks = self._create_chunks(text)

            logger.info(f"Processed CSV: {file_path.name}, extracted {len(chunks)} chunks")
            return chunks

        except Exception as e:
            logger.error(f"Error processing CSV file: {str(e)}")
            return []

    def process_docx(self, file_path: Path) -> List[Dict[str, Any]]:
        """
        Process DOCX file and extract text.

        Args:
            file_path: Path to DOCX file

        Returns:
            List of text chunks
        """
        try:
            doc = DocxDocument(file_path)

            # Extract text from paragraphs
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

            # Extract text from tables
            for table in doc.tables:
                table_text = self._format_docx_table(table)
                text += "\n\n" + table_text

            # Create chunks
            chunks = self._create_chunks(text)

            logger.info(f"Processed DOCX: {file_path.name}, extracted {len(chunks)} chunks")
            return chunks

        except Exception as e:
            logger.error(f"Error processing DOCX file: {str(e)}")
            return []

    def _create_chunks(
        self,
        text: str,
        page_number: Optional[int] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Dict[str, Any]]:
        """
        Create overlapping chunks from text.

        Args:
            text: Text to chunk
            page_number: Optional page number
            metadata: Optional metadata

        Returns:
            List of chunks with metadata
        """
        chunks = []
        text_length = len(text)

        for i in range(0, text_length, self.chunk_size - self.chunk_overlap):
            chunk_text = text[i:i + self.chunk_size]

            if chunk_text.strip():
                chunk = {
                    "text": chunk_text.strip(),
                    "page_number": page_number,
                    "metadata": metadata or {}
                }
                chunks.append(chunk)

        return chunks

    def _format_tables(self, tables: List[List[List[str]]]) -> str:
        """
        Format PDF tables as text.

        Args:
            tables: List of tables from pdfplumber

        Returns:
            Formatted table text
        """
        formatted = []

        for table in tables:
            if table:
                # Convert to DataFrame for better formatting
                df = pd.DataFrame(table[1:], columns=table[0] if table[0] else None)
                formatted.append(df.to_string(index=False))

        return "\n\n".join(formatted)

    def _format_docx_table(self, table) -> str:
        """
        Format DOCX table as text.

        Args:
            table: DOCX table object

        Returns:
            Formatted table text
        """
        data = []
        for row in table.rows:
            data.append([cell.text for cell in row.cells])

        if data:
            df = pd.DataFrame(data[1:], columns=data[0] if data[0] else None)
            return df.to_string(index=False)

        return ""

    def cleanup_file(self, file_path: Path) -> bool:
        """
        Delete temporary file after processing.

        Args:
            file_path: Path to file to delete

        Returns:
            True if deleted successfully
        """
        try:
            if file_path.exists():
                file_path.unlink()
                logger.info(f"Cleaned up file: {file_path.name}")
                return True
            return False

        except Exception as e:
            logger.error(f"Error cleaning up file {file_path}: {str(e)}")
            return False


# Global document processor instance
document_processor = DocumentProcessor()
